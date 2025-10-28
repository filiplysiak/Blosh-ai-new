Read the pdf and extract data from the right columns
import PyPDF2
import pandas as pd

# Read the PDF file
pdf_path = "data/Erm Fashion Branche Dames Rapportage W - Week 39 2025.pdf"

with open(pdf_path, 'rb') as file:
    pdf_reader = PyPDF2.PdfReader(file)
    
    # Get number of pages
    num_pages = len(pdf_reader.pages)
    print(f"Number of pages: {num_pages}")
    
    # Extract text from all pages
    full_text = ""
    for page_num in range(num_pages):
        page = pdf_reader.pages[page_num]
        full_text += page.extract_text()
    
    print("\n--- First 1000 characters ---")
    print(full_text[:1000])

import pdfplumber
import pandas as pd
import re

def extract_tables_from_pdf(pdf_path):
    """
    Extract two tables from the ERM Fashion PDF:
    1. Summary table (Totaal seizoen)
    2. Brand details table (Seizoen per merk)
    """
    
    with pdfplumber.open(pdf_path) as pdf:
        # Initialize variables
        summary_data = None
        brand_rows = []
        
        for page_num, page in enumerate(pdf.pages, 1):
            print(f"Processing page {page_num}...")
            
            # Extract tables from the page
            tables = page.extract_tables()
            
            for table in tables:
                if not table or len(table) == 0:
                    continue
                
                # Check if this is the summary table (page 1)
                if page_num == 1 and any('Totaal seizoen' in str(cell) for row in table for cell in row if cell):
                    # Find the row with actual data (after headers)
                    for row in table:
                        if row and len(row) > 5:
                            # Check if row contains numeric data
                            if any(re.search(r'\d+\.\d+', str(cell)) for cell in row if cell):
                                summary_data = row
                                break
                
                # Check if this is a brand table
                if any('Merk' in str(cell) for row in table for cell in row if cell):
                    # Find header row
                    header_idx = None
                    for i, row in enumerate(table):
                        if row and 'Merk' in str(row[0]):
                            header_idx = i
                            break
                    
                    if header_idx is not None:
                        # Extract data rows (skip header)
                        for row in table[header_idx + 1:]:
                            if row and row[0] and row[0].strip():  # Has brand name
                                # Skip if it's another header row
                                if 'Merk' not in str(row[0]) and 'Omzet' not in str(row[0]):
                                    brand_rows.append(row)
        
        # Create summary dataframe
        summary_columns = [
            'Omzet index ond', 'Omzet index grp', '%Dvk ond', '%Dvk grp',
            'Rent ond', 'Rent grp', 'Marge ond', 'Marge grp', 'OS ond', 'OS grp'
        ]
        
        if summary_data:
            # Clean and filter the summary data
            clean_summary = [cell for cell in summary_data if cell and cell.strip()]
            df_summary = pd.DataFrame([clean_summary[:10]], columns=summary_columns[:len(clean_summary)])
        else:
            df_summary = pd.DataFrame(columns=summary_columns)
        
        # Create brand dataframe
        brand_columns = [
            'Merk', 'Omzet index ond', 'Omzet index grp', '%Dvk ond', '%Dvk grp',
            'Rent ond', 'Rent grp', 'Aand grp', 'Marge grp', 'OS grp'
        ]
        
        df_brands = pd.DataFrame(brand_rows, columns=brand_columns[:len(brand_rows[0])] if brand_rows else brand_columns)
        
        # Clean up the dataframes
        for df in [df_summary, df_brands]:
            # Remove completely empty rows
            df.dropna(how='all', inplace=True)
        
        return df_summary, df_brands


# PDF path
pdf_path = "data/Erm Fashion Branche Dames Rapportage W - Week 39 2025.pdf"

# Extract tables
df_summary, df_brands = extract_tables_from_pdf(pdf_path)

display(df_summary)
display(df_brands)

df_brands.to_excel("brands_table.xlsx", index=False)
df_summary.to_excel("summary_table.xlsx", index=False)

import pandas as pd
from docx import Document
import re

def safe_numeric_convert(series, strip_percent=False):
    """Safely convert a series to numeric, handling percentages and mixed types"""
    if strip_percent:
        # Convert to string first, then strip %, then to numeric
        return pd.to_numeric(series.astype(str).str.rstrip('%'), errors='coerce')
    else:
        return pd.to_numeric(series, errors='coerce')

def extract_week_from_filename(pdf_path):
    """Extract week number from PDF filename"""
    # Look for pattern like "Week 39" or "W39" or similar
    match = re.search(r'[Ww]eek?\s*(\d+)', pdf_path)
    if match:
        return int(match.group(1))
    return 39  # Default fallback

def get_brand_data(df_brands, brand_name):
    """Get data for a specific brand from the dataframe"""
    brand_row = df_brands[df_brands['Merk'].str.upper() == brand_name.upper()]
    if len(brand_row) == 0:
        return {
            'omzet_index': 'N/A',
            'dvk': 'N/A',
            'rent': 'N/A',
            'marge': 'N/A',
            'os': 'N/A'
        }
    
    row = brand_row.iloc[0]
    return {
        'omzet_index': str(row['Omzet index grp']),
        'dvk': str(row['%Dvk grp']),
        'rent': str(row['Rent grp']),
        'marge': str(row['Marge grp']),
        'os': str(row['OS grp'])
    }

def create_top10_text(df, column, metric_name):
    """Create top 10 list as formatted text"""
    df_sorted = df.nlargest(10, column)
    text = ""
    for idx, (_, row) in enumerate(df_sorted.iterrows(), 1):
        col_value = row[column]
        text += f"{idx}. {row['Merk']} - {col_value}\n"
    return text.strip()

def analyze_volume_brands(df):
    """Analyze volume brands"""
    try:
        df = df.copy()
        df['Aand grp'] = safe_numeric_convert(df['Aand grp'], strip_percent=True)
        df['OS grp'] = safe_numeric_convert(df['OS grp'])
        
        high_volume = df[df['Aand grp'] > 1.0]
        
        if len(high_volume) == 0:
            return "Volumemerken analyse: Geen merken met aandeel > 1% gevonden in de data."
        
        text = "Volumemerken (Aandeel > 1%):\n\n"
        for _, row in high_volume.nlargest(10, 'Aand grp').iterrows():
            text += f"• {row['Merk']}: {row['Aand grp']:.1f}% aandeel, OS: {row['OS grp']}, Marge: {row['Marge grp']}\n"
        
        text += "\nDeze merken zijn de ruggengraat van de omzet en zorgen voor stabiele volumes."
        return text.strip()
    except Exception as e:
        return f"Volumemerken analyse: Data niet volledig beschikbaar. Error: {str(e)}"

def analyze_niche_brands(df):
    """Analyze niche brands with high margin"""
    try:
        df = df.copy()
        df['Aand grp'] = safe_numeric_convert(df['Aand grp'], strip_percent=True)
        df['Marge grp'] = safe_numeric_convert(df['Marge grp'], strip_percent=True)
        
        niche = df[(df['Aand grp'] < 1.0) & (df['Marge grp'] > 56)]
        
        if len(niche) == 0:
            return "Nichemerken analyse: Geen merken met laag aandeel (<1%) en hoge marge (>56%) gevonden."
        
        text = "Nichemerken (Laag aandeel < 1%, maar hoge marge > 56%):\n\n"
        for _, row in niche.nlargest(10, 'Marge grp').iterrows():
            text += f"• {row['Merk']}: {row['Aand grp']:.1f}% aandeel, Marge: {row['Marge grp']:.2f}%\n"
        
        text += "\nDeze merken bieden hoogwaardige winstmarges ondanks kleinere volumes."
        return text.strip()
    except Exception as e:
        return f"Nichemerken analyse: Data niet volledig beschikbaar. Error: {str(e)}"

def analyze_high_margin_brands(df):
    """Analyze brands with high margins"""
    try:
        df = df.copy()
        df['Marge grp'] = safe_numeric_convert(df['Marge grp'], strip_percent=True)
        high_margin = df[df['Marge grp'] > 56]
        
        if len(high_margin) == 0:
            return "Hoge marge analyse: Geen merken met marge > 56% gevonden."
        
        text = "Merken met hoge marges (>56%):\n\n"
        for _, row in high_margin.nlargest(10, 'Marge grp').iterrows():
            text += f"• {row['Merk']}: {row['Marge grp']:.2f}% marge, Rent: {row['Rent grp']}, OS: {row['OS grp']}\n"
        
        text += "\n→ Deze merken zijn winstgevend en strategisch belangrijk voor de totale marge."
        return text.strip()
    except Exception as e:
        return f"Hoge marge analyse: Data niet volledig beschikbaar. Error: {str(e)}"

def analyze_low_margin_brands(df):
    """Analyze brands with low margins"""
    try:
        df = df.copy()
        df['Marge grp'] = safe_numeric_convert(df['Marge grp'], strip_percent=True)
        low_margin = df[df['Marge grp'] < 48]
        
        if len(low_margin) == 0:
            return "Lage marge analyse: Geen merken met marge < 48% gevonden (goed nieuws!)."
        
        text = "Merken met lage marges (<48%):\n\n"
        for _, row in low_margin.nsmallest(10, 'Marge grp').iterrows():
            text += f"• {row['Merk']}: {row['Marge grp']:.2f}% marge, OS: {row['OS grp']}, Aandeel: {row['Aand grp']}\n"
        
        text += "\n→ Let op: deze merken kunnen volume brengen maar drukken de winstgevendheid."
        return text.strip()
    except Exception as e:
        return f"Lage marge analyse: Data niet volledig beschikbaar. Error: {str(e)}"

def create_brand_clustering(df):
    """Create brand clustering analysis"""
    try:
        df = df.copy()
        df['Marge grp'] = safe_numeric_convert(df['Marge grp'], strip_percent=True)
        df['Rent grp'] = safe_numeric_convert(df['Rent grp'])
        df['OS grp'] = safe_numeric_convert(df['OS grp'])
        
        text = "Brand Clustering:\n\n"
        
        # Winstmotors
        winners = df[(df['Marge grp'] > 55) & (df['Rent grp'] > 500)]
        text += "WINSTMOTORS (Hoge marge + hoog rendement):\n"
        if len(winners) > 0:
            for _, row in winners.head(5).iterrows():
                text += f"  • {row['Merk']}\n"
        else:
            text += "  Geen merken gevonden in deze categorie\n"
        
        # Exit kandidaten
        text += "\nEXIT KANDIDATEN (Lage OS + lage omzet):\n"
        exit_brands = df[(df['OS grp'] < 2.5) & (df['Omzet index grp'] < 1.0)]
        if len(exit_brands) > 0:
            for _, row in exit_brands.head(5).iterrows():
                text += f"  • {row['Merk']}\n"
        else:
            text += "  Geen merken gevonden in deze categorie\n"
        
        return text.strip()
    except Exception as e:
        return f"Brand clustering analyse: Data niet volledig beschikbaar. Error: {str(e)}"

def create_summary(freebird_data, gem_groep_data, competitors, df_numeric):
    """Create comprehensive Dutch summary based on performance"""
    try:
        omzet_fb = float(freebird_data['omzet_index'])
        omzet_gem = float(gem_groep_data['omzet_index'])
        
        # Strip % if present
        marge_fb_str = str(freebird_data['marge']).rstrip('%')
        marge_gem_str = str(gem_groep_data['marge']).rstrip('%')
        marge_fb = float(marge_fb_str)
        marge_gem = float(marge_gem_str)
        
        os_fb = float(freebird_data['os'])
        os_gem = float(gem_groep_data['os'])
        
        dvk_fb_str = str(freebird_data['dvk']).rstrip('%')
        dvk_gem_str = str(gem_groep_data['dvk']).rstrip('%')
        dvk_fb = float(dvk_fb_str)
        dvk_gem = float(dvk_gem_str)
        
        rent_fb = float(freebird_data['rent'])
        rent_gem = float(gem_groep_data['rent'])
    except Exception as e:
        return f"SAMENVATTING: Onvoldoende data beschikbaar voor volledige FREEBIRD analyse. Error: {str(e)}"
    
    text = "SAMENVATTING - FREEBIRD Performance Week-tot-Datum:\n\n"
    
    # === OVERALL PERFORMANCE ===
    text += "=== OVERALL PERFORMANCE ===\n\n"
    
    # Omzet analyse
    omzet_gap = ((omzet_fb / omzet_gem) - 1) * 100
    if omzet_fb > omzet_gem:
        text += f"[+] OMZET: FREEBIRD presteert {omzet_gap:.1f}% BOVEN het groepsgemiddelde\n"
        text += f"  Index: {omzet_fb:.2f} vs. {omzet_gem:.2f} (groep)\n"
        if omzet_gap > 15:
            text += f"  → Uitstekende performance! FREEBIRD is een duidelijke trekker in het assortiment.\n"
        else:
            text += f"  → Goede performance, maar er is ruimte voor verdere groei.\n"
    else:
        text += f"[-] OMZET: FREEBIRD blijft {abs(omzet_gap):.1f}% ACHTER op het groepsgemiddelde\n"
        text += f"  Index: {omzet_fb:.2f} vs. {omzet_gem:.2f} (groep)\n"
        text += f"  → Actie vereist: analyseer waarom het merk onderpresteert en pas strategie aan.\n"
    
    text += "\n"
    
    # Marge analyse
    marge_gap = marge_fb - marge_gem
    if marge_fb > marge_gem:
        text += f"[+] MARGE: {marge_fb:.2f}% - {marge_gap:.1f}pp BOVEN groepsgemiddelde ({marge_gem:.2f}%)\n"
        text += f"  → Gezonde winstgevendheid! FREEBIRD draagt sterk bij aan de totale marge.\n"
    else:
        text += f"[-] MARGE: {marge_fb:.2f}% - {abs(marge_gap):.1f}pp ONDER groepsgemiddelde ({marge_gem:.2f}%)\n"
        text += f"  → Focus op margin recovery: herzie inkoop- en pricing strategie.\n"
    
    text += "\n"
    
    # Doorverkoop analyse
    dvk_gap = dvk_fb - dvk_gem
    if dvk_fb > dvk_gem:
        text += f"[+] DOORVERKOOP: {dvk_fb:.1f}% - {dvk_gap:.1f}pp BOVEN gemiddelde ({dvk_gem:.1f}%)\n"
        if dvk_fb > 40:
            text += f"  → Sterke sell-through! Het assortiment sluit goed aan bij klantvraag.\n"
        else:
            text += f"  → Acceptabele doorverkoop, maar blijf monitoren voor seizoenseinde.\n"
    else:
        text += f"[-] DOORVERKOOP: {dvk_fb:.1f}% - {abs(dvk_gap):.1f}pp ONDER gemiddelde ({dvk_gem:.1f}%)\n"
        if dvk_fb < 30:
            text += f"  → Zorgwekkend! Overweeg promotie-acties om voorraad te activeren.\n"
        else:
            text += f"  → Onder gemiddelde maar nog acceptabel. Monitor nauwlettend.\n"
    
    text += "\n"
    
    # OS (voorraadrotatie) analyse
    os_gap = ((os_fb / os_gem) - 1) * 100
    if os_fb > os_gem:
        text += f"[+] VOORRAADROTATIE (OS): {os_fb:.2f} - {os_gap:.1f}% SNELLER dan gemiddelde ({os_gem:.2f})\n"
        if os_fb > 4:
            text += f"  → Excellente rotatie! Voorraad wordt snel verkocht zonder overstock.\n"
        else:
            text += f"  → Goede rotatie, voorraad beweegt gezond door.\n"
    else:
        text += f"[-] VOORRAADROTATIE (OS): {os_fb:.2f} - {abs(os_gap):.1f}% LANGZAMER dan gemiddelde ({os_gem:.2f})\n"
        if os_fb < 2:
            text += f"  → Kritiek! Voorraad staat te lang. Risico op deadstock.\n"
        else:
            text += f"  → Rotatie kan beter. Evalueer assortiment en presentatie.\n"
    
    text += "\n"
    
    # Rentabiliteit analyse
    rent_gap = ((rent_fb / rent_gem) - 1) * 100
    text += f"• RENTABILITEIT: EUR {rent_fb:.2f} per stuk (groep: EUR {rent_gem:.2f})\n"
    if rent_fb > rent_gem:
        text += f"  → {rent_gap:.1f}% winstgevender per verkocht item dan gemiddeld.\n"
    else:
        text += f"  → {abs(rent_gap):.1f}% minder winstgevend per item. Focus op margin verbetering.\n"
    
    # === CONCURRENTIE POSITIE ===
    text += "\n\n=== POSITIE T.O.V. DIRECTE CONCURRENTEN ===\n\n"
    
    try:
        # Vergelijk met key concurrenten
        josh_v_omzet = float(competitors['JOSH V']['omzet_index'])
        fabienne_omzet = float(competitors['FABIENNE CHAPOT']['omzet_index'])
        harper_omzet = float(competitors['HARPER & YVE']['omzet_index'])
        
        text += "Omzetindex vergelijking:\n"
        text += f"  • FREEBIRD: {omzet_fb:.2f}\n"
        text += f"  • Josh V: {josh_v_omzet:.2f}\n"
        text += f"  • Fabienne Chapot: {fabienne_omzet:.2f}\n"
        text += f"  • Harper & Yve: {harper_omzet:.2f}\n\n"
        
        # Bepaal positie
        competitors_list = [
            ('Josh V', josh_v_omzet),
            ('Fabienne Chapot', fabienne_omzet),
            ('Harper & Yve', harper_omzet),
            ('FREEBIRD', omzet_fb)
        ]
        competitors_list.sort(key=lambda x: x[1], reverse=True)
        freebird_rank = [i for i, (name, _) in enumerate(competitors_list, 1) if name == 'FREEBIRD'][0]
        
        text += f"→ FREEBIRD staat op positie {freebird_rank} van 4 in deze peer group.\n"
        
        if freebird_rank == 1:
            text += "  Uitstekend! FREEBIRD is de sterkste performer in het segment.\n"
        elif freebird_rank == 2:
            text += "  Sterke positie, maar er is een benchmark om naar te streven.\n"
        else:
            text += "  Er is significant groeipotentieel door te leren van betere performers.\n"
    except Exception as e:
        text += f"Concurrentie vergelijking: Data niet volledig beschikbaar.\n"
    
    # === STRATEGISCHE AANBEVELINGEN ===
    text += "\n\n=== STRATEGISCHE AANBEVELINGEN ===\n\n"
    
    # Genereer specifieke aanbevelingen op basis van performance
    recommendations = []
    
    if omzet_fb < omzet_gem:
        recommendations.append("[OMZET] Verhoog zichtbaarheid op winkelvloer (prime locations, meer facings)")
        recommendations.append("[OMZET] Analyseer welke collectie-items wel/niet presteren en pas inkoop aan")
    
    if marge_fb < marge_gem:
        recommendations.append("[MARGE] Heronderhandel inkoopprijzen met leverancier")
        recommendations.append("[MARGE] Minimaliseer kortingsacties; focus op full-price verkoop")
    
    if dvk_fb < 35:
        recommendations.append("[DOORVERKOOP] Start targeted promotions om doorverkoop te stimuleren")
        recommendations.append("[DOORVERKOOP] Cross-sell met beter presterende merken")
    
    if os_fb < 3:
        recommendations.append("[VOORRAAD] Kritisch review van slow-movers; overweeg retour of markdown")
        recommendations.append("[VOORRAAD] Reduceer inkoopvolume volgende periode tot OS > 3.0")
    
    if not recommendations:
        recommendations.append("[OK] Continue huidige strategie - performance is gezond")
        recommendations.append("[OK] Monitor wekelijks voor behoud van deze positie")
        recommendations.append("[OK] Zoek incrementele verbetermogelijkheden in assortiment mix")
    
    for rec in recommendations:
        text += f"{rec}\n"
    
    # === ACTIEPUNTEN KOMENDE WEEK ===
    text += "\n\n=== ACTIEPUNTEN KOMENDE WEEK ===\n\n"
    text += "[ ] Wekelijkse review van doorverkoop% per collectie-onderdeel\n"
    text += "[ ] Bespreek FREEBIRD performance in team-overleg\n"
    text += "[ ] Vergelijk in-store presentatie met best-performing concurrenten\n"
    text += "[ ] Evalueer of huidige voorraadniveau aansluit bij verkoopsnelheid\n"
    text += "[ ] Plan inkoop volgende seizoen op basis van deze week's inzichten\n"
    
    # === OVERALL CONCLUSIE ===
    text += "\n\n=== OVERALL CONCLUSIE ===\n\n"
    
    # Tel aantal positieve vs negatieve indicatoren
    positive_score = sum([
        omzet_fb > omzet_gem,
        marge_fb > marge_gem,
        dvk_fb > dvk_gem,
        os_fb > os_gem,
        rent_fb > rent_gem
    ])
    
    if positive_score >= 4:
        text += "[STERK] FREEBIRD is een STERK MERK in het portfolio met gezonde KPI's over de linie.\n"
        text += "De focus ligt op het behouden en verder uitbouwen van deze positie.\n"
    elif positive_score >= 3:
        text += "[ACCEPTABEL] FREEBIRD presteert ACCEPTABEL met meer sterke dan zwakke punten.\n"
        text += "Met gerichte verbeteringen kan dit merk een top-performer worden.\n"
    elif positive_score >= 2:
        text += "[GEMIXED] FREEBIRD heeft een GEMIXTE PERFORMANCE met zowel kansen als zorgen.\n"
        text += "Prioriteer de zwakke KPI's voor verbetering in de komende weken.\n"
    else:
        text += "[KRITIEK] FREEBIRD ONDERPRESTEERT op meerdere vlakken en vereist DIRECTE ACTIE.\n"
        text += "Een strategische review van het merk is noodzakelijk: herpositioneren of afbouwen?\n"
    
    return text.strip()

def analyze_season_comparison(df_summary, freebird_data):
    """Analyze season comparison for FREEBIRD using year-over-year data"""
    try:
        # FREEBIRD specifieke data
        omzet_fb_dj = float(freebird_data['omzet_index'])
        dvk_fb_dj = float(str(freebird_data['dvk']).rstrip('%'))
        marge_fb_dj = float(str(freebird_data['marge']).rstrip('%'))
        
        # Groep totaal data
        omzet_grp_dj = float(df_summary['Omzet index grp'].iloc[0])
        dvk_grp_dj = float(df_summary['%Dvk grp'].iloc[0])
        marge_grp_dj = float(str(df_summary['Marge grp'].iloc[0]).rstrip('%'))
        
        text = "Seizoensvergelijking (Jaar-op-Jaar):\n\n"
        
        text += "FREEBIRD Performance:\n"
        text += f"• Omzetindex: {omzet_fb_dj:.2f}\n"
        text += f"• Doorverkoop: {dvk_fb_dj:.1f}%\n"
        text += f"• Marge: {marge_fb_dj:.1f}%\n"
        
        text += "\nTotale Groep Performance:\n"
        text += f"• Omzetindex: {omzet_grp_dj:.2f} (FREEBIRD is {((omzet_fb_dj/omzet_grp_dj)-1)*100:.1f}% vs groep)\n"
        text += f"• Doorverkoop: {dvk_grp_dj:.1f}% (FREEBIRD is {dvk_fb_dj - dvk_grp_dj:.1f}pp vs groep)\n"
        text += f"• Marge: {marge_grp_dj:.1f}% (FREEBIRD is {marge_fb_dj - marge_grp_dj:.1f}pp vs groep)\n"
        
        text += "\nCONCLUSIE: "
        if omzet_fb_dj > omzet_grp_dj and marge_fb_dj > marge_grp_dj:
            text += "FREEBIRD presteert boven groepsgemiddelde met goede marges - voortzetten strategie."
        elif omzet_fb_dj > omzet_grp_dj:
            text += "FREEBIRD heeft goede omzet maar marges kunnen beter."
        else:
            text += "FREEBIRD heeft groeipotentieel - focus op omzetverbetering en margeverbetering."
        
        return text.strip()
    except Exception as e:
        return f"Seizoensvergelijking: Historische vergelijkingsdata niet beschikbaar. Huidige week data wordt getoond zonder vergelijking."

def analyze_doorverkoop_trend(df):
    """Analyze sell-through (DvK) trend"""
    try:
        df = df.copy()
        df['%Dvk grp'] = safe_numeric_convert(df['%Dvk grp'])
        
        avg_dvk = df['%Dvk grp'].mean()
        
        text = "Doorverkooptrend (%DvK) Analyse:\n\n"
        text += f"Gemiddelde doorverkoop: {avg_dvk:.1f}%\n\n"
        
        # High performers
        high_dvk = df[df['%Dvk grp'] > 45]
        text += "TOPPERS (>45% doorverkoop):\n"
        if len(high_dvk) > 0:
            for _, row in high_dvk.nlargest(5, '%Dvk grp').iterrows():
                text += f"  • {row['Merk']}: {row['%Dvk grp']:.1f}% - {row['OS grp']} OS\n"
        else:
            text += "  Geen merken met >45% doorverkoop\n"
        
        # Low performers
        text += "\nACHTERBLIJVERS (<30% doorverkoop):\n"
        low_dvk = df[df['%Dvk grp'] < 30]
        if len(low_dvk) > 0:
            for _, row in low_dvk.nsmallest(5, '%Dvk grp').iterrows():
                text += f"  • {row['Merk']}: {row['%Dvk grp']:.1f}% - {row['OS grp']} OS\n"
        else:
            text += "  Geen merken met <30% doorverkoop\n"
        
        text += "\nADVIES:\n"
        text += "• Merken met <30% doorverkoop: evalueer pricing en promotie-acties\n"
        text += "• Merken met >45% doorverkoop: verzeker voldoende voorraad voor restant seizoen\n"
        
        return text.strip()
    except Exception as e:
        return f"Doorverkooptrend analyse: Data niet volledig beschikbaar. Error: {str(e)}"

def analyze_voorraad_risico(df):
    """Analyze inventory risk based on OS and DvK"""
    try:
        df = df.copy()
        df['%Dvk grp'] = safe_numeric_convert(df['%Dvk grp'])
        df['OS grp'] = safe_numeric_convert(df['OS grp'])
        
        text = "Voorraadrisico Analyse (OS + %DvK):\n\n"
        
        # High risk: Low OS + Low DvK
        high_risk = df[(df['OS grp'] < 2.5) & (df['%Dvk grp'] < 30)]
        text += "HOOG RISICO (Trage verkoop + lage doorverkoop):\n"
        if len(high_risk) > 0:
            for _, row in high_risk.head(5).iterrows():
                text += f"  • {row['Merk']}: OS {row['OS grp']:.2f}, DvK {row['%Dvk grp']:.1f}%\n"
            text += "  → Actie: Overweeg afprijzing of retourneren naar leverancier\n"
        else:
            text += "  Geen merken met hoog voorraadrisico geidentificeerd\n"
        
        # Medium risk: One metric is low
        text += "\nGEMIDDELD RISICO:\n"
        medium_risk = df[((df['OS grp'] < 2.5) & (df['%Dvk grp'] >= 30)) | 
                         ((df['OS grp'] >= 2.5) & (df['%Dvk grp'] < 30))]
        if len(medium_risk) > 0:
            for _, row in medium_risk.head(5).iterrows():
                text += f"  • {row['Merk']}: OS {row['OS grp']:.2f}, DvK {row['%Dvk grp']:.1f}%\n"
            text += "  → Actie: Monitor wekelijks en pas strategie aan indien nodig\n"
        else:
            text += "  Geen merken met gemiddeld risico\n"
        
        # Low risk: High OS + High DvK
        text += "\nLAAG RISICO (Snelle verkoop + hoge doorverkoop):\n"
        low_risk = df[(df['OS grp'] > 4) & (df['%Dvk grp'] > 40)]
        if len(low_risk) > 0:
            for _, row in low_risk.head(5).iterrows():
                text += f"  • {row['Merk']}: OS {row['OS grp']:.2f}, DvK {row['%Dvk grp']:.1f}%\n"
            text += "  → Deze merken presteren uitstekend\n"
        else:
            text += "  Beperkte merken met optimale voorraadrotatie\n"
        
        return text.strip()
    except Exception as e:
        return f"Voorraadrisico analyse: Data niet volledig beschikbaar. Error: {str(e)}"

def replace_text_in_paragraph(paragraph, replacements):
    """Replace placeholder text in a paragraph while preserving formatting"""
    # First, try to merge runs to handle split placeholders
    full_text = paragraph.text
    
    # Check if any placeholder exists in this paragraph
    needs_replacement = False
    for placeholder in replacements.keys():
        if placeholder in full_text:
            needs_replacement = True
            break
    
    if not needs_replacement:
        return
    
    # Perform replacements
    for placeholder, value in replacements.items():
        full_text = full_text.replace(placeholder, str(value))
    
    # Clear the paragraph and add new text
    # Save the first run's formatting
    first_run_format = None
    if paragraph.runs:
        first_run_format = paragraph.runs[0]
    
    # Clear all runs
    for run in paragraph.runs:
        run.text = ''
    
    # Re-add text with original formatting if possible
    if first_run_format is not None:
        paragraph.runs[0].text = full_text
    else:
        paragraph.add_run(full_text)

def replace_text_in_table(table, replacements):
    """Replace placeholder text in a table while preserving formatting"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph, replacements)

def fill_template(template_path, df_summary, df_brands, pdf_path=None, week_number=None, year=2025):
    """Fill the template DOCX with actual values from dataframes"""
    
    # Extract week number from PDF path if not provided
    if week_number is None and pdf_path:
        week_number = extract_week_from_filename(pdf_path)
    elif week_number is None:
        week_number = 39  # Default fallback
    
    # Load the template
    doc = Document(template_path)
    
    # Prepare numeric dataframe
    df_numeric = df_brands.copy()
    df_numeric['Omzet index grp'] = safe_numeric_convert(df_numeric['Omzet index grp'])
    df_numeric['%Dvk grp'] = safe_numeric_convert(df_numeric['%Dvk grp'])
    df_numeric['Rent grp'] = safe_numeric_convert(df_numeric['Rent grp'])
    df_numeric['Marge grp'] = safe_numeric_convert(df_numeric['Marge grp'], strip_percent=True)
    df_numeric['OS grp'] = safe_numeric_convert(df_numeric['OS grp'])
    
    # Get competitor data
    competitors = {
        'FREEBIRD': get_brand_data(df_brands, 'FREEBIRD'),
        'FABIENNE CHAPOT': get_brand_data(df_brands, 'FABIENNE CHAPOT'),
        'HARPER & YVE': get_brand_data(df_brands, 'HARPER & YVE'),
        'JOSH V': get_brand_data(df_brands, 'JOSH V'),
        'POM AMSTERDAM': get_brand_data(df_brands, 'POM AMSTERDAM'),
        'AAIKO': get_brand_data(df_brands, 'AAIKO')
    }
    
    # Calculate group average
    gem_groep = {
        'omzet_index': f"{df_numeric['Omzet index grp'].mean():.2f}",
        'dvk': f"{df_numeric['%Dvk grp'].mean():.2f}",
        'rent': f"{df_numeric['Rent grp'].mean():.2f}",
        'marge': f"{df_numeric['Marge grp'].mean():.2f}%",
        'os': f"{df_numeric['OS grp'].mean():.2f}"
    }
    
    # Create replacements dictionary
    replacements = {
        '{{weeknummer}}': str(week_number),
        '{{jaar}}': str(year),
        '{{vorige_week}}': str(week_number - 1),
        
        # FREEBIRD
        '{{omzet_index_freebird}}': competitors['FREEBIRD']['omzet_index'],
        '{{dvk_freebird}}': competitors['FREEBIRD']['dvk'],
        '{{rent_freebird}}': competitors['FREEBIRD']['rent'],
        '{{marge_freebird}}': competitors['FREEBIRD']['marge'],
        '{{os_freebird}}': competitors['FREEBIRD']['os'],
        
        # Fabienne Chapot
        '{{omzet_index_fabienne_chapot}}': competitors['FABIENNE CHAPOT']['omzet_index'],
        '{{dvk_fabienne_chapot}}': competitors['FABIENNE CHAPOT']['dvk'],
        '{{rent_fabienne_chapot}}': competitors['FABIENNE CHAPOT']['rent'],
        '{{marge_fabienne_chapot}}': competitors['FABIENNE CHAPOT']['marge'],
        '{{os_fabienne_chapot}}': competitors['FABIENNE CHAPOT']['os'],
        
        # Harper & Yve
        '{{omzet_index_harper_&_yve}}': competitors['HARPER & YVE']['omzet_index'],
        '{{dvk_harper_&_yve}}': competitors['HARPER & YVE']['dvk'],
        '{{rent_harper_&_yve}}': competitors['HARPER & YVE']['rent'],
        '{{marge_harper_&_yve}}': competitors['HARPER & YVE']['marge'],
        '{{os_harper_&_yve}}': competitors['HARPER & YVE']['os'],
        
        # Josh V
        '{{omzet_index_josh_v}}': competitors['JOSH V']['omzet_index'],
        '{{dvk_josh_v}}': competitors['JOSH V']['dvk'],
        '{{rent_josh_v}}': competitors['JOSH V']['rent'],
        '{{marge_josh_v}}': competitors['JOSH V']['marge'],
        '{{os_josh_v}}': competitors['JOSH V']['os'],
        
        # POM Amsterdam
        '{{omzet_index_pom_amsterdam}}': competitors['POM AMSTERDAM']['omzet_index'],
        '{{dvk_pom_amsterdam}}': competitors['POM AMSTERDAM']['dvk'],
        '{{rent_pom_amsterdam}}': competitors['POM AMSTERDAM']['rent'],
        '{{marge_pom_amsterdam}}': competitors['POM AMSTERDAM']['marge'],
        '{{os_pom_amsterdam}}': competitors['POM AMSTERDAM']['os'],
        
        # Aaiko
        '{{omzet_index_aaiko}}': competitors['AAIKO']['omzet_index'],
        '{{dvk_aaiko}}': competitors['AAIKO']['dvk'],
        '{{rent_aaiko}}': competitors['AAIKO']['rent'],
        '{{marge_aaiko}}': competitors['AAIKO']['marge'],
        '{{os_aaiko}}': competitors['AAIKO']['os'],
        
        # Gem. groep
        '{{omzet_index_gem._groep}}': gem_groep['omzet_index'],
        '{{dvk_gem._groep}}': gem_groep['dvk'],
        '{{rent_gem._groep}}': gem_groep['rent'],
        '{{marge_gem._groep}}': gem_groep['marge'],
        '{{os_gem._groep}}': gem_groep['os'],
        
        # Analyses
        '{{top10_omzetindex}}': create_top10_text(df_numeric, 'Omzet index grp', 'omzetindex'),
        '{{top10_rent}}': create_top10_text(df_numeric, 'Rent grp', 'rentabiliteit'),
        '{{top10_os}}': create_top10_text(df_numeric, 'OS grp', 'OS rotatie'),
        '{{analyse_volumemerken}}': analyze_volume_brands(df_numeric.copy()),
        '{{analyse_nichemerken}}': analyze_niche_brands(df_numeric.copy()),
        '{{merken_hoge_marge}}': analyze_high_margin_brands(df_numeric.copy()),
        '{{merken_lage_marge}}': analyze_low_margin_brands(df_numeric.copy()),
        '{{brand_clustering}}': create_brand_clustering(df_numeric.copy()),
        '{{samenvatting}}': create_summary(competitors['FREEBIRD'], gem_groep, competitors, df_numeric),
        
        # New analyses
        '{{seizoensvergelijking}}': analyze_season_comparison(df_summary, competitors['FREEBIRD']),
        '{{doorverkooptrend}}': analyze_doorverkoop_trend(df_numeric.copy()),
        '{{voorraadrisico}}': analyze_voorraad_risico(df_numeric.copy())
    }
    
    # Replace text in all paragraphs
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)
    
    # Replace text in all tables
    for table in doc.tables:
        replace_text_in_table(table, replacements)
    
    # Save the filled document
    output_filename = f'Analyse ERM Brancherapportage Week {week_number} {year}.docx'
    doc.save(output_filename)
    
    print(f"✓ Rapport gegenereerd: {output_filename}")
    return output_filename

# Template path and PDF path
template_path = "data/template_blosh.docx"
pdf_path = "Erm Fashion Branche Dames Rapportage W - Week 39 2025.pdf"  # Update this to your actual PDF filename

# Fill the template with data (week number will be extracted automatically from pdf_path)
output_file = fill_template(template_path, df_summary, df_brands, pdf_path=pdf_path, year=2025)
print(f"\n✓ Template succesvol gevuld en opgeslagen als: {output_file}")